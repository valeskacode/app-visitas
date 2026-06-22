import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import os

# --- CONFIGURACIÓN Y ESTILOS ---
st.set_page_config(page_title="Auditoría Caja Arequipa", layout="centered")

# --- AUTENTICACIÓN ---
if "logged_in" not in st.session_state: st.session_state.logged_in = False

def login():
    st.title("🔐 Acceso de Auditor")
    user = st.text_input("Usuario")
    pwd = st.text_input("Contraseña", type="password")
    if st.button("Ingresar"):
        if user == "auditor" and pwd == "caja2026":
            st.session_state.logged_in = True
            st.rerun()
        else:
            st.error("Credenciales inválidas")

if not st.session_state.logged_in:
    login()
else:
    # --- LÓGICA DE NEGOCIO ---
    if "form_data" not in st.session_state:
        st.session_state.form_data = {}

    st.markdown("<h2 style='color:#8B0000;'>📋 Formato de Visita</h2>", unsafe_allow_html=True)
    paso = st.select_slider("Sección", options=["1. Datos", "2. Riesgos", "3. Campo", "4. Negocio", "5. Cierre"])

    with st.form("main_form"):
        if paso == "1. Datos":
            st.session_state.form_data['tituclie'] = st.text_input("Titular", st.session_state.form_data.get('tituclie', ''))
            st.session_state.form_data['cuentacliente'] = st.text_input("Cuenta", st.session_state.form_data.get('cuentacliente', ''))
        elif paso == "2. Riesgos":
            st.session_state.form_data['deudadirecta'] = st.number_input("Deuda Directa", value=st.session_state.form_data.get('deudadirecta', 0.0))
        elif paso == "3. Campo":
            st.session_state.form_data['direccion'] = st.text_input("Dirección", st.session_state.form_data.get('direccion', ''))
            st.session_state.form_data['comentarios'] = st.text_area("Comentarios", st.session_state.form_data.get('comentarios', ''))
        elif paso == "4. Negocio":
            st.session_state.form_data['actividadprincipal'] = st.text_input("Actividad Principal", st.session_state.form_data.get('actividadprincipal', ''))
        elif paso == "5. Cierre":
            st.session_state.form_data['iniau'] = st.text_input("Hecho por", st.session_state.form_data.get('iniau', ''))
            submitted = st.form_submit_button("Guardar en memoria")

    # --- FUNCIÓN DE REEMPLAZO EN PLANTILLA ---
    def generar_documento(datos):
        if not os.path.exists("plantilla_visita.docx"):
            return None
        doc = Document("plantilla_visita.docx")
        
        def reemplazar(objeto):
            for key, value in datos.items():
                target = f"[{key}]"
                if target in objeto.text:
                    objeto.text = objeto.text.replace(target, str(value))
        
        for p in doc.paragraphs: reemplazar(p)
        for t in doc.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs: reemplazar(p)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # --- DESCARGA ---
    if st.button("📥 Generar Documento Final"):
        archivo = generar_documento(st.session_state.form_data)
        if archivo:
            st.download_button("Descargar Informe.docx", data=archivo, 
                               file_name="Informe_Auditoria_Final.docx", 
                               mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        else:
            st.error("Error: No se encontró la plantilla 'plantilla_visita.docx'.")

    if st.sidebar.button("Cerrar Sesión"):
        st.session_state.logged_in = False
        st.rerun()
